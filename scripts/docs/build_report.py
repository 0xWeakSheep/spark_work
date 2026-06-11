#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "report" / "第X组_电影数据离线分析系统_课程总结报告.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(85, 85, 85)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    run.font.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("电影数据离线分析系统课程总结报告")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(42)
    r = p.add_run("成都信息工程大学\nSpark大数据实验报告")
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = DARK_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("基于 Spark 的电影数据离线分析系统")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = BLUE

    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(4.0), Cm(10.5)]
    labels = ["学院", "专业", "班级", "小组", "姓名", "学号"]
    values = ["________________", "________________", "________________", "第 X 组", "所有小组成员", "所有小组成员"]
    for row, label, value in zip(table.rows, labels, values):
        row.cells[0].width = widths[0]
        row.cells[1].width = widths[1]
        set_cell_text(row.cells[0], label, True)
        set_cell_text(row.cells[1], value)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    doc.add_heading("目录", level=1)
    items = [
        "1. 项目原理",
        "2. 平台搭建过程",
        "3. 项目总体设计",
        "4. 项目数据库设计",
        "5. 数据处理过程详细设计",
        "6. 程序实现",
        "7. 问题及解决",
        "8. 结论及结果展示",
        "9. 总结",
        "附录 A. 运行命令与验收清单",
    ]
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)
    doc.add_page_break()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_text(cell, header, True)
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            set_cell_text(cell, value)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(10.5)
    p = cell.add_paragraph()
    p.add_run(body)
    doc.add_paragraph()


def add_screenshot_placeholder(doc: Document, caption: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FBFCFE")
    cell.height = Cm(5)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"截图占位：{caption}")
    r.font.color.rgb = MUTED
    r.font.size = Pt(11)
    doc.add_paragraph(caption).alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_paragraphs(doc: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(22)
        p.add_run(text)


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "0F172A")
    p = cell.paragraphs[0]
    run = p.add_run(code)
    run.font.name = "Menlo"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(226, 232, 240)
    doc.add_paragraph()


def add_section_page(doc: Document, title: str, subtitle: str, paragraphs: list[str]) -> None:
    doc.add_heading(title, level=1)
    add_callout(doc, "本节目标", subtitle)
    add_paragraphs(doc, paragraphs)


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_footer(doc)
    add_cover(doc)
    add_toc(doc)

    add_section_page(
        doc,
        "1. 项目原理",
        "说明 Spark、HDFS、MySQL 与 Web 可视化在本项目中的协同关系。",
        [
            "本项目以 MovieLens 电影评分数据为对象，构建从数据存储、离线计算、结果入库到可视化展示的完整链路。HDFS 负责存放原始数据文件，Spark 负责对评分、电影和用户数据进行分布式处理，MySQL 负责保存结构化结果，Flask 与 ECharts 负责将分析结果转换为可交互图表。",
            "Spark 的核心能力体现在 RDD、Spark SQL 与 Streaming 三个方面。RDD 用于完成评分文件的基础转换、聚合和排序；Spark SQL 用于完成跨表关联、电影类型拆分和按性别聚合；Streaming 用于模拟评分文件持续到达时的增量统计过程。",
            "系统采用单机伪分布式架构，不追求生产集群规模，而是完整体现课程要求中的 Linux、Hadoop、Spark、MySQL 与 Web 技术整合。HDFS 副本数设置为 1，能够减少本机资源占用，同时保持 Hadoop 文件系统接口和 Spark 读取方式一致。",
        ],
    )
    add_table(
        doc,
        ["组件", "项目中的职责", "对应考核点"],
        [
            ["Hadoop HDFS", "保存 Movies、Ratings、Tags、Users 数据文件", "数据上传至 HDFS"],
            ["Spark RDD", "计算平均评分最高前 20 部电影", "RDD 过滤、聚合、排序"],
            ["Spark SQL", "关联 MySQL 表并统计男女关注电影类型", "Spark SQL 访问 MySQL"],
            ["Spark Streaming", "监听新增评分批次并写入趋势表", "实时折线图变化"],
            ["Flask + ECharts", "提供接口和图表展示", "Web 整合与可视化"],
        ],
    )
    add_screenshot_placeholder(doc, "系统总体架构图")
    doc.add_page_break()

    add_section_page(
        doc,
        "1.1 Spark 分布式基本原理",
        "解释 Spark 任务在 Driver、Executor 和数据分区上的执行方式。",
        [
            "Spark 应用由 Driver 负责生成逻辑计划和调度任务，由 Executor 负责在分区上执行具体计算。评分数据被划分为多个分区后，每个分区独立解析 CSV 行并提取 movieId、rating 等字段，后续通过 reduceByKey 聚合评分总和与评分次数。",
            "在本系统中，电影数据规模约百万级评分记录，单机 local[2] 足以完成计算。虽然运行在单机环境，但代码仍使用 Spark 的分布式算子和 DataFrame API，后续迁移到多节点集群时主要调整 master 和资源参数即可。",
            "RDD 适合展示 Spark 底层转换过程，DataFrame 和 SQL 适合表达结构化关联与聚合，Streaming 则体现持续数据处理模型。三者组合能够覆盖课程考核要求的主要知识点。",
        ],
    )
    add_code_block(
        doc,
        "ratings.reduceByKey(lambda left, right: (left[0] + right[0], left[1] + right[1]))\n"
        "averages = rating_stats.mapValues(lambda item: (item[0] / item[1], item[1]))\n"
        "top20 = averages.join(movies).sortBy(lambda item: (-item[1][0][0], item[1][1])).take(20)",
    )
    doc.add_page_break()

    add_section_page(
        doc,
        "1.2 Spark 核心数据抽象",
        "对应 RDD、DataFrame、临时视图和流式 DataFrame 的使用。",
        [
            "RDD 是弹性分布式数据集，适合表达 map、filter、reduceByKey、join 等基础算子。本项目的 Top 20 电影评分榜使用 RDD 实现，保留了从文本文件到键值对、从聚合到排序的完整过程。",
            "DataFrame 是带 Schema 的分布式数据结构，适合 SQL 查询和跨数据源整合。本项目通过 JDBC 读取 MySQL 中的 movies、ratings、users 表，使用 DataFrame join 和 explode 拆分电影类型，最后通过 SQL 计算每个性别对各类型的关注数量。",
            "Streaming 处理的是持续到达的数据流。系统把评分数据拆分为小批次文件并追加到 HDFS 目录，Spark 以流式 DataFrame 读取目录变化，对每个批次计算评分数量、平均分、最高分和最低分。",
        ],
    )
    add_table(
        doc,
        ["抽象", "输入", "输出", "用途"],
        [
            ["RDD", "HDFS CSV 文本", "top_movies 表", "平均评分排序"],
            ["DataFrame", "MySQL 原始表", "聚合结果集合", "结构化关联分析"],
            ["SQL View", "DataFrame 临时视图", "SQL 聚合结果", "男女类型统计"],
            ["Streaming DataFrame", "HDFS 新增文件", "stream_rating_trend 表", "实时趋势"],
        ],
    )
    doc.add_page_break()

    add_section_page(
        doc,
        "2. 平台搭建过程",
        "说明 Docker 一体化环境、端口、依赖版本和启动流程。",
        [
            "考虑到 macOS 上直接安装旧版本 Hadoop、Spark 和 Java 的兼容性成本较高，本项目使用 Docker Compose 搭建一体化环境。MySQL、Hadoop/Spark 和 Flask 分别运行在独立容器中，通过 Compose 网络互相访问。",
            "Hadoop/Spark 容器基于 OpenJDK 11 镜像，下载 Hadoop 3.3.6 与 Spark 3.5.1，并安装 PyMySQL 作为 Spark 作业写入 MySQL 的客户端依赖。MySQL 使用 8.4 官方镜像，Web 服务使用 Python 3.11 slim 镜像。",
            "本机只需要 Docker Desktop 即可运行，不需要在 macOS 系统中安装 Hadoop 或 Spark。Docker Desktop 建议分配 8GB 内存，HDFS 副本数设置为 1，适合 16GB 内存的 MacBook Air。",
        ],
    )
    add_table(
        doc,
        ["服务", "端口", "说明"],
        [
            ["movie-mysql", "3307 -> 3306", "MySQL 数据库，保存原始表与分析结果"],
            ["movie-spark-hadoop", "9870, 4040", "HDFS NameNode UI 与 Spark UI"],
            ["movie-web", "5000", "Flask 可视化仪表盘"],
        ],
    )
    add_code_block(doc, "docker compose up -d --build\ndocker compose exec spark-hadoop bash scripts/setup_all.sh")
    add_screenshot_placeholder(doc, "Docker Compose 服务启动截图")
    doc.add_page_break()

    add_section_page(
        doc,
        "2.1 数据集说明",
        "说明 MovieLens 数据文件、字段和本项目使用方式。",
        [
            "当前数据包解压后包含 Movies.csv、Ratings.csv、Tags.csv、Users.dat 和 README.txt。Movies.csv 提供电影 ID、片名和类型列表；Ratings.csv 提供用户评分记录；Tags.csv 提供用户标签；Users.dat 提供用户性别、年龄、职业和邮编。",
            "清洗后数据量为：电影 45,843 条，评分 1,048,575 条，标签 753,155 条，用户 6,040 条。清洗过程主要处理编码容错、空值过滤、评分范围过滤以及 Users.dat 到 Users.csv 的结构化转换。",
            "由于 Users.dat 的用户范围为 1 到 6040，而 Tags.csv 覆盖更大的 userId 范围，男女关注统计选择 Ratings.csv + Movies.csv + Users.dat 作为主统计口径，按电影 genres 拆分为电影标签/类型，保证性别关联数据充足且逻辑稳定。",
        ],
    )
    add_table(
        doc,
        ["文件", "主要字段", "用途"],
        [
            ["Movies.csv", "movieId, title, genres", "电影标题与类型维度"],
            ["Ratings.csv", "userId, movieId, rating, timestamp", "评分聚合与用户兴趣"],
            ["Users.dat", "UserID, Gender, Age, Occupation, Zip-code", "性别维度统计"],
            ["Tags.csv", "userId, movieId, tag, timestamp", "补充标签数据入库"],
        ],
    )
    doc.add_page_break()

    add_section_page(
        doc,
        "3. 项目总体设计",
        "给出系统分层、数据流和功能模块设计。",
        [
            "系统分为数据层、计算层、存储层、服务层和展示层。数据层负责保存清洗后的 CSV 文件和流式批次文件；计算层由 Spark RDD、Spark SQL 和 Streaming 作业组成；存储层由 MySQL 原始表与结果表组成；服务层由 Flask 提供 API；展示层由 ECharts 绘制图表。",
            "用户访问首页时，页面会调用四个 API：概览统计、Top 20 电影、男女类型统计和流式趋势。后端只读取 MySQL 结果表，不直接执行 Spark 作业，因此页面响应稳定，计算任务和展示任务解耦。",
            "批处理作业在初始化或手动运行时执行，Streaming 作业在演示时启动。这样的设计便于课程答辩：可以先展示离线结果，再运行流式脚本观察折线图变化。",
        ],
    )
    add_screenshot_placeholder(doc, "Web 仪表盘首页截图")
    doc.add_page_break()

    add_section_page(
        doc,
        "3.1 功能结构设计",
        "按照考核要求拆分系统功能。",
        [
            "功能一是 RDD 评分榜。系统从 HDFS 读取电影和评分数据，过滤表头和异常行，以 movieId 为键聚合评分总和和评分次数，计算平均分后与电影标题关联，输出平均评分最高的 20 部电影。",
            "功能二是 Spark SQL 性别类型统计。系统先将原始数据导入 MySQL，再通过 Spark JDBC 读取结构化表，关联评分、电影和用户信息，拆分电影类型后按 gender、genre 聚合，最终生成男女关注最多的电影类型。",
            "功能三是 Streaming 趋势统计。系统把 Ratings.csv 拆分为多个批次文件，定时上传到 HDFS 流式目录，Spark 监听目录新增文件并按批次计算评分记录数和平均分，结果写入 MySQL 供页面折线图刷新。",
        ],
    )
    add_table(
        doc,
        ["功能", "输入", "处理方式", "展示方式"],
        [
            ["Top 20 电影", "HDFS Movies/Ratings", "RDD 聚合排序", "横向柱状图 + 表格"],
            ["男女类型统计", "MySQL movies/ratings/users", "Spark SQL join + explode", "分组柱状图"],
            ["实时评分趋势", "HDFS stream input", "Structured Streaming foreachBatch", "折线图"],
        ],
    )
    doc.add_page_break()

    add_section_page(
        doc,
        "4. 项目数据库设计",
        "说明 MySQL 原始表和结果表。",
        [
            "数据库 movie_analysis 包含四张原始表和四张结果表。原始表用于保留清洗后的输入数据，结果表用于存放 Spark 计算后的可视化数据。这样的设计可以避免 Web 页面直接扫描大数据文件，也便于检查计算结果。",
            "movies 表以 movie_id 为主键，ratings 表对 user_id 和 movie_id 建索引，users 表以 user_id 为主键并对 gender 建索引，tags 表对 user_id、movie_id 和 tag 建索引。结果表根据查询模式设置主键或唯一键。",
            "stream_rating_trend 表使用 batch_id 唯一约束，Streaming 作业重复处理相同批次时会执行 upsert，避免因为检查点或演示重跑导致重复数据。",
        ],
    )
    add_table(
        doc,
        ["表名", "类型", "说明"],
        [
            ["movies", "原始表", "电影 ID、标题、类型"],
            ["ratings", "原始表", "用户评分记录"],
            ["users", "原始表", "用户性别、年龄、职业"],
            ["tags", "原始表", "用户自定义标签"],
            ["top_movies", "结果表", "评分最高前 20 部电影"],
            ["gender_genre_attention", "结果表", "男女电影类型关注数量"],
            ["gender_top_genres", "结果表", "男女最关注类型"],
            ["stream_rating_trend", "结果表", "流式评分趋势"],
        ],
    )
    add_screenshot_placeholder(doc, "MySQL 表结构截图")
    doc.add_page_break()

    add_section_page(
        doc,
        "5. 数据处理过程详细设计",
        "从解压、清洗、上传、导入到分析的端到端流程。",
        [
            "第一步是解压 data/archive/moviedata-latest.rar 到 data/raw/moviedata-latest。第二步是运行 scripts/data/prepare_data.py 对 CSV 和 DAT 文件进行清洗，输出到 data/processed。第三步是 scripts/storage/setup_hdfs.sh 将四个主要数据文件上传到 HDFS 的 /movie/input 目录。",
            "第四步是 load_mysql.py 将清洗后的数据批量导入 MySQL 原始表。第五步是依次执行 rdd_top_movies.py 和 sql_gender_genres.py，分别生成 top_movies、gender_genre_attention 和 gender_top_genres。第六步是在演示时执行 run_streaming_demo.sh，生成 stream_rating_trend。",
            "整个流程采用脚本化实现，既便于复现实验，也便于课程报告中展示每一步的命令、输出和截图。",
        ],
    )
    add_code_block(
        doc,
        "bash scripts/data/extract_data.sh\n"
        "python3 scripts/data/prepare_data.py\n"
        "bash scripts/storage/setup_hdfs.sh\n"
        "python3 scripts/storage/load_mysql.py\n"
        "spark-submit src/spark/rdd_top_movies.py\n"
        "spark-submit src/spark/sql_gender_genres.py",
    )
    add_screenshot_placeholder(doc, "HDFS /movie/input 文件列表截图")
    doc.add_page_break()

    sections = [
        (
            "5.1 RDD 处理流程",
            "RDD 作业从 HDFS 文本文件开始，使用 csv.reader 在 mapPartitions 中解析，避免电影标题中逗号导致字段错位。评分记录转换为 movieId -> (rating, 1)，经过 reduceByKey 得到评分总和与次数，再计算平均分。",
            "RDD 结果与 movies RDD 按 movieId join 后得到标题信息，最后按照平均分降序、评分次数降序和电影标题升序排序，取前 20 写入 MySQL 的 top_movies 表。页面横向柱状图和明细表均读取该结果表。",
        ),
        (
            "5.2 Spark SQL 处理流程",
            "Spark SQL 作业通过 JDBC 读取 MySQL 原始表，使用 ratings join users join movies 得到带性别与电影类型的评分记录。genres 字段以竖线分隔，作业使用 explode(split(genres, '\\\\|')) 将一部电影的多个类型拆为多行。",
            "拆分后的数据按 gender 和 genre 分组计数，代表不同性别用户对电影类型的关注程度。随后使用窗口函数 row_number 选出每个性别关注人数最多的类型，分别写入 gender_genre_attention 与 gender_top_genres。",
        ),
        (
            "5.3 Streaming 处理流程",
            "流式演示脚本先把 Ratings.csv 拆为 8 个小文件，每个文件 1000 行，再按固定间隔上传到 HDFS 的 /movie/stream/input 目录。Spark Structured Streaming 使用 maxFilesPerTrigger=1 控制每次处理一个批次文件。",
            "每个批次通过 foreachBatch 进入 write_batch 函数，计算 rating_count、avg_rating、min_rating 和 max_rating，然后写入 stream_rating_trend。Web 页面每 6 秒轮询一次 API，因此新增批次会表现为折线图不断增长。",
        ),
        (
            "5.4 Web 展示流程",
            "Flask 提供 /api/summary、/api/top-movies、/api/gender-genres 和 /api/stream-trend 四个接口。前端页面首次加载时并行请求全部接口，并在点击刷新按钮或定时器触发时更新图表。",
            "ECharts 图表包括 Top 20 横向柱状图、男女类型分组柱状图和评分趋势折线图。图表配色采用蓝色、青绿色和琥珀色组合，既保持数据仪表盘的专业感，也避免页面只由单一颜色构成。",
        ),
    ]
    for heading, p1, p2 in sections:
        doc.add_heading(heading, level=2)
        add_paragraphs(doc, [p1, p2])
        add_screenshot_placeholder(doc, f"{heading} 截图")
        doc.add_page_break()

    add_section_page(
        doc,
        "6. 程序实现",
        "列出关键模块和核心代码职责。",
        [
            "项目代码按职责划分为 src、scripts、infra、config、data、docs 和 artifacts。src 负责 Spark 计算与 Web 应用，scripts 负责流程编排，infra 负责容器配置，config 负责数据库初始化，data 保存数据，docs 保存报告与答辩材料，artifacts 保存生成物和日志。",
            "实现中尽量避免把计算逻辑写入 Web 层。Web 层只查询 MySQL 结果表，这样既能降低页面请求延迟，也能清晰体现离线分析系统的分层思想。",
            "Spark 作业的参数通过环境变量读取，默认值适合 Docker Compose 环境。后续如果迁移到独立 Hadoop/Spark 集群，只需要调整 HDFS 路径、MySQL 地址和 Spark master 参数。",
        ],
    )
    add_table(
        doc,
        ["模块", "代表文件", "职责"],
        [
            ["容器环境", "docker-compose.yml", "编排 MySQL、Hadoop/Spark、Flask 服务"],
            ["数据准备", "scripts/data/prepare_data.py", "清洗和规范化数据"],
            ["数据导入", "scripts/storage/load_mysql.py", "批量写入 MySQL 原始表"],
            ["RDD 作业", "src/spark/rdd_top_movies.py", "评分榜计算"],
            ["SQL 作业", "src/spark/sql_gender_genres.py", "男女类型关注统计"],
            ["Streaming 作业", "src/spark/streaming_ratings.py", "实时评分趋势计算"],
            ["Web 展示", "src/web/app.py", "API 与模板渲染"],
        ],
    )
    doc.add_page_break()

    implementation_pages = [
        ("6.1 Docker Compose 实现", "Compose 文件定义三个服务。MySQL 负责数据库，spark-hadoop 容器负责 HDFS 与 Spark 作业运行，web 容器负责 Flask 页面。各服务共享同一个 Compose 网络，Spark 作业通过 mysql 主机名访问数据库。"),
        ("6.2 数据清洗实现", "prepare_data.py 使用 Python csv 模块解析 CSV，避免手写字符串切分导致标题中的逗号破坏字段。脚本还将 Users.dat 转换为 Users.csv，方便 MySQL 批量导入。"),
        ("6.3 MySQL 导入实现", "load_mysql.py 使用 PyMySQL executemany 分批提交，每批 5000 行。导入前 TRUNCATE 原始表，保证重复执行 setup_all.sh 时不会累积脏数据。"),
        ("6.4 RDD 作业实现", "rdd_top_movies.py 使用 SparkContext 读取 HDFS CSV，mapPartitions 解析、reduceByKey 聚合、join 标题、sortBy 排序，最后将前 20 写入 top_movies。"),
        ("6.5 SQL 作业实现", "sql_gender_genres.py 使用 SparkSession 通过 JDBC 读取 MySQL 表，使用 DataFrame join 和 Spark SQL 窗口函数完成统计。"),
        ("6.6 Streaming 作业实现", "streaming_ratings.py 通过 readStream 监听 HDFS 目录，对每个新增文件批次计算评分指标，并用 ON DUPLICATE KEY UPDATE 写入 MySQL。"),
        ("6.7 Flask 接口实现", "app.py 将数据库查询封装为 API，返回 JSON 格式数据。异常时返回 ok=false，前端可以在控制台定位错误。"),
        ("6.8 ECharts 页面实现", "dashboard.js 初始化三个图表并负责数据刷新。页面使用响应式 CSS 网格，能适配桌面和较窄屏幕。"),
    ]
    for heading, body in implementation_pages:
        doc.add_heading(heading, level=2)
        add_paragraphs(doc, [body, "本模块的实现遵循职责单一原则，输入、处理和输出路径明确，便于在答辩中按模块展示代码和运行结果。"])
        add_screenshot_placeholder(doc, f"{heading} 代码或运行截图")
        doc.add_page_break()

    add_section_page(
        doc,
        "7. 问题及解决",
        "记录实现过程中遇到的关键问题和解决方案。",
        [
            "问题一是 macOS 本机 Java 版本与 Spark/Hadoop 旧版本兼容性不稳定。解决方式是使用 Docker 容器固定 OpenJDK 11、Hadoop 3.3.6 和 Spark 3.5.1，避免污染本机环境。",
            "问题二是数据文件编码存在少量异常字节，直接按 UTF-8 严格解析会中断。解决方式是在清洗脚本中使用 UTF-8-SIG 并开启 errors='replace'，保留可用字段并跳过无效记录。",
            "问题三是 Users.dat 和 Tags.csv 的用户范围不一致。解决方式是将男女关注统计的主口径改为 Ratings.csv + Movies.csv + Users.dat，按电影 genres 统计关注类型，Tags.csv 作为补充原始表导入。",
            "问题四是 Streaming 演示需要可观察变化。解决方式是将 Ratings.csv 拆分为多个小批次并定时上传到 HDFS，由页面定时轮询 MySQL 趋势表。",
        ],
    )
    add_table(
        doc,
        ["问题", "原因", "解决方法"],
        [
            ["Java 兼容性", "本机 Java 23 与大数据组件组合不稳定", "Docker 固定 OpenJDK 11"],
            ["编码错误", "部分电影标题或标签存在异常字节", "清洗脚本容错读取"],
            ["用户范围不一致", "Tags.csv 用户范围大于 Users.dat", "性别统计使用评分数据口径"],
            ["Streaming 可视化", "静态文件无法体现实时变化", "批次文件模拟流式输入"],
        ],
    )
    doc.add_page_break()

    add_section_page(
        doc,
        "8. 结论及结果展示",
        "总结运行后应展示的核心结果。",
        [
            "系统运行完成后，首页展示四个概览指标、Top 20 电影评分榜、男女电影类型关注柱状图和实时评分趋势折线图。Top 20 图表体现 RDD 批处理结果，男女类型图体现 Spark SQL 关联分析结果，趋势图体现 Streaming 结果。",
            "报告截图建议包含 Docker 容器列表、HDFS 文件列表、MySQL 表数据、Spark 作业运行输出、Web 首页全图、Top 20 图表、男女类型图和 Streaming 折线图变化前后对比。",
            "通过这些截图可以证明系统完成了从数据上传、Spark 分析、MySQL 保存到页面展示的完整链路。",
        ],
    )
    for caption in [
        "Docker 容器运行状态",
        "HDFS 输入目录文件",
        "MySQL top_movies 表结果",
        "Top 20 电影评分图",
        "男女关注电影类型柱状图",
        "Streaming 折线图变化截图",
    ]:
        add_screenshot_placeholder(doc, caption)
        doc.add_page_break()

    add_section_page(
        doc,
        "9. 总结",
        "归纳项目收获、技术理解和后续改进方向。",
        [
            "本项目完成了 Spark 大数据技术课程要求中的核心环节：HDFS 数据存储、Spark RDD 离线计算、Spark SQL 访问 MySQL、Spark Streaming 实时模拟、MySQL 结果保存和 Web 可视化展示。",
            "通过项目实现，可以更清楚地理解大数据系统并不是单一计算程序，而是由数据采集、存储、计算、服务和展示共同组成的工程系统。脚本化和容器化能够提升实验复现效率，也能减少环境差异带来的问题。",
            "后续可以继续扩展推荐算法、电影类型筛选、时间范围筛选、用户年龄段分析和更多实时指标，使系统从课程展示进一步发展为更完整的电影数据分析平台。",
        ],
    )
    doc.add_page_break()

    add_section_page(
        doc,
        "附录 A. 运行命令与验收清单",
        "列出复现实验所需命令。",
        [
            "首次运行使用 docker compose up -d --build 构建并启动服务，然后在 spark-hadoop 容器中执行 scripts/setup_all.sh 完成数据准备、HDFS 上传、MySQL 导入和批处理作业。",
            "流式演示使用 scripts/streaming/run_streaming_demo.sh，脚本会自动清理旧输入、创建批次文件、启动 Streaming 作业并向 HDFS 追加文件。演示过程中打开 Web 页面即可观察趋势图变化。",
        ],
    )
    add_code_block(
        doc,
        "docker compose up -d --build\n"
        "docker compose exec spark-hadoop bash scripts/setup_all.sh\n"
        "docker compose exec spark-hadoop hdfs dfs -ls /movie/input\n"
        "docker compose exec spark-hadoop bash scripts/streaming/run_streaming_demo.sh",
    )
    add_table(
        doc,
        ["验收项", "预期结果"],
        [
            ["HDFS 文件", "/movie/input 下包含 Movies.csv、Ratings.csv、Tags.csv、Users.dat"],
            ["MySQL 原始表", "movies、ratings、users、tags 均有数据"],
            ["RDD 结果", "top_movies 有 20 条记录"],
            ["SQL 结果", "gender_genre_attention 和 gender_top_genres 有统计结果"],
            ["Streaming 结果", "stream_rating_trend 随批次增加记录"],
            ["Web 页面", "三个核心图表正常显示"],
        ],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
